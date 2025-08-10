from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def annotate_docx(input_path, issues, output_path):
    """
    Adds inline issue notes for each detected problem
    and appends a summary section at the end.
    """
    doc = Document(input_path)

    for para in doc.paragraphs:
        for issue in issues:
            doc_type = issue.get("document", "")
            if doc_type and (doc_type in para.text or doc_type in input_path):
                run = para.add_run(
                    f"\n[ISSUE: {issue.get('issue', '')}] "
                    f"[Law: {issue.get('law_citation', '')}]"
                )
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True

   
    doc.add_page_break()

    para = doc.add_paragraph("Compliance Review Summary")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.runs[0]
    run.font.size = Pt(16)
    run.bold = True

 
    for issue in issues:
        summary_text = (
            f"{issue.get('document', 'N/A')} | {issue.get('issue', '')} "
            f"(Severity: {issue.get('severity', '')})\n"
            f"Law: {issue.get('law_citation', '')}\n"
            f"Suggestion: {issue.get('suggestion', '')}"
        )
        para = doc.add_paragraph(summary_text)
        if para.runs:
            para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    doc.save(output_path)
    return output_path
